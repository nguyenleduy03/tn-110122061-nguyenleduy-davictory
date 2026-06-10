"""Word document (.docx) parser."""

import io

from models.import_models import ParseResult


class DocxParser:
    def parse(self, content: bytes, filename: str) -> ParseResult:
        result = ParseResult(success=True, filename=filename, file_type="docx")
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            result.raw_text = "\n\n".join(parts)
            result.text_length = len(result.raw_text)
        except Exception as e:
            result.success = False
            result.raw_text = str(e)
        return result
